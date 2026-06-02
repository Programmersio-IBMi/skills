#!/usr/bin/env python3
"""
Professional Markdown-to-Word (DOCX) converter with comprehensive formatting support.
Produces polished technical documents suitable for distribution.

Usage: python convert_md_to_docx.py input.md [output.docx]

Requirements: pip install python-docx requests
"""

import argparse
import base64
import io
import os
import re
import sys
from typing import List, Optional, Tuple, Dict, Any

try:
    import requests
except ImportError:
    requests = None

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx is not installed.")
    print("Install with: pip install python-docx")
    sys.exit(1)

# ── Color Palette ──────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1a, 0x36, 0x5d)       # Dark blue
COLOR_SECONDARY = RGBColor(0x2c, 0x52, 0x82)     # Medium blue
COLOR_ACCENT = RGBColor(0x31, 0x82, 0xce)        # Bright blue
COLOR_MUTED = RGBColor(0x71, 0x80, 0x96)         # Gray
COLOR_TABLE_HEADER = RGBColor(0x2b, 0x6c, 0xb0)  # Table header blue
COLOR_CODE_BG = RGBColor(0xf7, 0xfa, 0xfc)       # Light gray background
COLOR_ALT_ROW = RGBColor(0xeb, 0xf4, 0xff)       # Alternate row blue
COLOR_BLOCKQUOTE = RGBColor(0x4a, 0x55, 0x68)    # Blockquote gray
COLOR_INLINE_CODE = RGBColor(0xc7, 0x25, 0x4e)   # Code red
COLOR_LINK = RGBColor(0x31, 0x82, 0xce)          # Link blue
COLOR_CODE_TEXT = RGBColor(0x2d, 0x37, 0x48)     # Code block text

CODE_BG_HEX = 'F7FAFC'
INLINE_CODE_BG_HEX = 'F1F5F9'
HEADER_ROW_HEX = '2B6CB0'
ALT_ROW_HEX = 'EBF4FF'

# ── Paragraph/Frame helpers ────────────────────────────────────────────


def set_cell_background(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_paragraph_shading(para, color_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)


def set_run_shading(run, color_hex: str):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    rPr.append(shd)


def add_horizontal_line(doc, color='CCCCCC'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_paragraph_borders(para, color='CCCCCC', size='4'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)


def add_left_border(para, color='3182CE', size='18'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_hyperlink(paragraph, url: str, text: str):
    """Insert a real hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '3182CE')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ── Inline formatting ──────────────────────────────────────────────────


def strip_markdown_formatting(text: str) -> str:
    """Remove markdown formatting from text, keeping the content."""
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\s)([^*\n]+?)(?<!\s)\*(?!\*)', r'\1', text)
    text = re.sub(r'`+([^`]+)`+', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


_XML_BAD = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')


def _xml_clean(s: str) -> str:
    return _XML_BAD.sub('', s) if s else s


def parse_inline_formatting(para, text: str):
    """Parse inline markdown and add formatted runs to the paragraph.

    Order:
      1. Code spans (greedy, protected first)
      2. Links
      3. Bold+italic ***text***
      4. Bold **text**
      5. Italic *text* / _text_
    """
    if not text:
        return

    code_spans: List[str] = []
    bi_spans: List[str] = []
    bold_spans: List[str] = []
    italic_spans: List[str] = []
    link_spans: List[Tuple[str, str]] = []

    def save_code(m):
        idx = len(code_spans)
        code_spans.append(m.group(1))
        return f'\x00C{idx}\x00'

    def save_link(m):
        idx = len(link_spans)
        link_spans.append((m.group(1), m.group(2)))
        return f'\x00L{idx}\x00'

    def save_bi(m):
        idx = len(bi_spans)
        bi_spans.append(m.group(1))
        return f'\x00X{idx}\x00'

    def save_bold(m):
        idx = len(bold_spans)
        bold_spans.append(m.group(1))
        return f'\x00B{idx}\x00'

    def save_italic(m):
        idx = len(italic_spans)
        italic_spans.append(m.group(1))
        return f'\x00I{idx}\x00'

    # Code spans (handle multi-tick fences first then single-tick)
    text = re.sub(r'``([^`]+)``', save_code, text)
    text = re.sub(r'`([^`\n]+)`', save_code, text)
    # Links
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', save_link, text)
    # Bold+italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', save_bi, text)
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', save_bold, text)
    # Italic — only when starred token is non-empty and not whitespace-bounded
    text = re.sub(r'(?<![A-Za-z0-9*])\*(?!\s)([^*\n]+?)(?<!\s)\*(?![A-Za-z0-9*])', save_italic, text)
    # Underscore italic (whole-word boundaries)
    text = re.sub(r'(?<![A-Za-z0-9_])_(?!\s)([^_\n]+?)(?<!\s)_(?![A-Za-z0-9_])', save_italic, text)

    # Decode placeholders — recursively, since bold/italic spans may contain
    # nested code-span or link placeholders that were saved earlier.
    placeholder_re = re.compile(r'(\x00[CLXBI]\d+\x00)')

    def render(segment: str, bold: bool = False, italic: bool = False):
        if not segment:
            return
        for part in placeholder_re.split(segment):
            if not part:
                continue
            m = re.match(r'\x00([CLXBI])(\d+)\x00$', part)
            if not m:
                run = para.add_run(_xml_clean(part))
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                continue
            kind, idx_s = m.group(1), int(m.group(2))
            if kind == 'C':
                run = para.add_run(_xml_clean(code_spans[idx_s]))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_INLINE_CODE
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
            elif kind == 'L':
                ltext, lurl = link_spans[idx_s]
                add_hyperlink(para, lurl, ltext)
            elif kind == 'X':
                render(bi_spans[idx_s], bold=True, italic=True)
            elif kind == 'B':
                render(bold_spans[idx_s], bold=True, italic=italic)
            elif kind == 'I':
                render(italic_spans[idx_s], bold=bold, italic=True)

    render(text)


# ── Tables ─────────────────────────────────────────────────────────────


def parse_table(lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
    rows = []
    idx = start_idx
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line.startswith('|'):
            break
        if re.match(r'^[\|\-\:\s]+$', line):
            idx += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        idx += 1
    return rows, idx


def add_formatted_table(doc, rows: List[List[str]]):
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    num_rows = len(rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Disable fixed layout so columns size to content
    tblPr = table._tbl.tblPr
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'autofit')

    for i, row_data in enumerate(rows):
        # Pad short rows
        if len(row_data) < num_cols:
            row_data = row_data + [''] * (num_cols - len(row_data))
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            # Markdown <br> inside cell -> line break
            segments = re.split(r'<br\s*/?>', cell_text, flags=re.IGNORECASE)
            for seg_idx, segment in enumerate(segments):
                if seg_idx > 0:
                    para.add_run().add_break()
                parse_inline_formatting(para, segment.strip())

            if i == 0:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10.5)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_background(cell, HEADER_ROW_HEX)
            else:
                if i % 2 == 0:
                    set_cell_background(cell, ALT_ROW_HEX)
                plain_text = strip_markdown_formatting(cell_text).strip()
                if plain_text and re.match(r'^-?\d+(\.\d+)?%?$', plain_text):
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ── Mermaid diagrams ───────────────────────────────────────────────────


MERMAID_INK_URL = 'https://mermaid.ink/img/{b64}?type=png'


def add_mermaid_diagram(doc, code_lines: List[str]) -> bool:
    """Render a mermaid block as a PNG via mermaid.ink and embed it.

    Returns True on success. On any failure, returns False so the caller can
    fall back to rendering the diagram source as a plain code block.
    """
    if requests is None:
        return False
    src = '\n'.join(code_lines).strip()
    if not src:
        return False
    try:
        b64 = base64.urlsafe_b64encode(src.encode('utf-8')).decode('ascii')
        resp = requests.get(MERMAID_INK_URL.format(b64=b64), timeout=20)
        resp.raise_for_status()
        if not resp.content:
            return False
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run()
        run.add_picture(io.BytesIO(resp.content), width=Inches(6.0))
        return True
    except Exception as exc:
        sys.stderr.write(f"[warn] mermaid render failed, falling back to code block: {exc}\n")
        return False


# ── Code blocks ────────────────────────────────────────────────────────


def add_code_block(doc, lines: List[str], language: str = ''):
    if not lines:
        # Drop trailing blank lines
        return
    # Trim leading/trailing blank lines inside the fence
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if not lines:
        return

    # Optional language label as a small caption above
    if language:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(0)
        run = cap.add_run(language.upper())
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_MUTED
        run.bold = True

    # Each code line is its own paragraph (no spacing between them) so Word
    # renders the block compactly with proper monospace alignment.
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.4)
        # Preserve indentation by using a non-breaking representation
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_CODE_TEXT
        set_paragraph_shading(p, CODE_BG_HEX)

        # Add top border on first line, bottom on last for a visual frame
        first = idx == 0
        last = idx == len(lines) - 1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side, on in (('top', first), ('bottom', last), ('left', True), ('right', True)):
            if not on:
                continue
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '4')
            b.set(qn('w:color'), 'CBD5E0')
            pBdr.append(b)
        pPr.append(pBdr)


# ── Lists ──────────────────────────────────────────────────────────────


def _list_indent_units(leading_ws: str) -> int:
    """Return indent level by counting tabs or 2/4-space groups."""
    spaces = 0
    for ch in leading_ws:
        if ch == '\t':
            spaces += 4
        else:
            spaces += 1
    # Each indent step is 2 spaces (commonly 2 or 4); 4 spaces = level 2.
    return spaces // 2


UL_RE = re.compile(r'^(\s*)([-*+])\s+(.*)$')
OL_RE = re.compile(r'^(\s*)(\d+)\.\s+(.*)$')


def is_list_line(line: str) -> bool:
    return bool(UL_RE.match(line) or OL_RE.match(line))


def parse_list_block(lines: List[str], start_idx: int) -> Tuple[List[Dict], int]:
    """Parse a contiguous list block. Returns items with level/ordered/text/number."""
    items: List[Dict] = []
    idx = start_idx
    base_indent = None

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.rstrip()

        # Allow blank lines inside a list (they don't end it if next line is a list item too)
        if not stripped.strip():
            # Look ahead — if next non-blank is a list item, skip blank; else stop.
            j = idx + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and is_list_line(lines[j]):
                idx = j
                continue
            break

        m_ul = UL_RE.match(stripped)
        m_ol = OL_RE.match(stripped)
        if not (m_ul or m_ol):
            # Continuation line for previous item if indented
            if items and (raw.startswith(' ') or raw.startswith('\t')):
                items[-1]['text'] += ' ' + stripped.strip()
                idx += 1
                continue
            break

        if m_ul:
            ws, _bullet, text = m_ul.group(1), m_ul.group(2), m_ul.group(3)
            ordered = False
            number = None
        else:
            ws, num, text = m_ol.group(1), m_ol.group(2), m_ol.group(3)
            ordered = True
            number = num

        level_units = _list_indent_units(ws)
        if base_indent is None:
            base_indent = level_units
        level = max(0, level_units - base_indent)
        level = min(level, 4)

        items.append({
            'ordered': ordered,
            'level': level,
            'text': text.rstrip(),
            'number': number,
        })
        idx += 1

    return items, idx


def add_list_items(doc, items: List[Dict]):
    """Render list items with manual numbering / bullets.

    Manual numbering avoids Word's continued-numbering behavior across
    document sections (e.g. Section 6 starting at "15." because earlier
    lists already used the List Number style).

    Each contiguous ordered block restarts at 1 unless the markdown
    explicitly starts higher.
    """
    bullets_by_level = ['•', '◦', '▪', '▫']

    # Per-level counters that restart at the first ordered item of this block
    counters: Dict[int, int] = {}

    for idx, item in enumerate(items):
        text = item['text']
        level = min(max(item['level'], 0), 3)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
        # Hanging indent so wrapped lines align under the text, not the marker
        p.paragraph_format.first_line_indent = Cm(-0.55)

        if item['ordered']:
            # Restart counter for this level the first time we hit it.
            # If user explicitly numbered the first item (e.g. starts at 5),
            # honor that starting number.
            if level not in counters:
                try:
                    start = int(item['number']) if item['number'] else 1
                except (TypeError, ValueError):
                    start = 1
                counters[level] = start
            n = counters[level]
            counters[level] += 1
            marker = f'{n}.'
        else:
            # Bullets reset any deeper-level ordered counters
            counters.pop(level + 1, None)
            counters.pop(level + 2, None)
            marker = bullets_by_level[min(level, len(bullets_by_level) - 1)]

        marker_run = p.add_run(marker + '\t')
        marker_run.font.name = 'Calibri'

        # Set a tab stop so the marker and text align cleanly
        from docx.shared import Cm as _Cm
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(_Cm(0.6 + 0.5 * level + 0.55))

        parse_inline_formatting(p, text)


# ── Document defaults ──────────────────────────────────────────────────


def configure_document_defaults(doc: Document):
    """Configure tight, professional defaults for the whole document."""
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Normal paragraph: tight, professional defaults
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20

    # Heading styles — color + spacing
    heading_settings = [
        ('Heading 1', COLOR_PRIMARY, Pt(20), True, Pt(18), Pt(8)),
        ('Heading 2', COLOR_SECONDARY, Pt(16), True, Pt(14), Pt(6)),
        ('Heading 3', COLOR_ACCENT, Pt(13), True, Pt(10), Pt(4)),
        ('Heading 4', COLOR_MUTED, Pt(12), True, Pt(8), Pt(2)),
        ('Heading 5', COLOR_MUTED, Pt(11), True, Pt(6), Pt(2)),
        ('Heading 6', COLOR_MUTED, Pt(11), True, Pt(4), Pt(2)),
    ]
    for name, color, size, bold, sb, sa in heading_settings:
        if name in doc.styles:
            st = doc.styles[name]
            st.font.color.rgb = color
            st.font.size = size
            st.font.bold = bold
            st.font.name = 'Calibri'
            st.paragraph_format.space_before = sb
            st.paragraph_format.space_after = sa
            st.paragraph_format.keep_with_next = True

    # List styles: tighten spacing
    for sname in ('List Bullet', 'List Bullet 2', 'List Bullet 3',
                  'List Number', 'List Number 2', 'List Number 3'):
        if sname in doc.styles:
            st = doc.styles[sname]
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.space_after = Pt(2)
            st.paragraph_format.line_spacing = 1.15


# ── Block detection helpers ────────────────────────────────────────────


def is_table_start(lines: List[str], i: int) -> bool:
    if i >= len(lines) or not lines[i].lstrip().startswith('|'):
        return False
    # Need at least the header line plus a separator-ish next line to be a table.
    if i + 1 < len(lines):
        nxt = lines[i + 1].strip()
        if re.match(r'^\|?[\s\-\:\|]+\|?$', nxt) and '-' in nxt:
            return True
    return False


# ── Main converter ────────────────────────────────────────────────────


def convert_md_to_docx(md_file: str, docx_file: Optional[str] = None) -> str:
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            raw = f.read()
    except FileNotFoundError:
        print(f"Error: Input file '{md_file}' not found.")
        sys.exit(1)
    except UnicodeDecodeError:
        with open(md_file, 'r', encoding='latin-1') as f:
            raw = f.read()

    lines = raw.splitlines()

    if docx_file is None:
        docx_file = os.path.splitext(md_file)[0] + '.docx'

    doc = Document()
    core_props = doc.core_properties
    core_props.title = os.path.splitext(os.path.basename(md_file))[0]
    core_props.author = "iA by programmers.io"

    configure_document_defaults(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line: skip — paragraph spacing handles separation.
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if m:
                level = len(m.group(1))
                title = strip_markdown_formatting(m.group(2)).strip()
                # Trim trailing # tokens (ATX style)
                title = re.sub(r'\s+#+\s*$', '', title)
                p = doc.add_heading(level=level)
                # Inline formatting allowed inside heading text
                parse_inline_formatting(p, title)
                if level == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                i += 1
                continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            add_horizontal_line(doc)
            i += 1
            continue

        # Fenced code block
        if stripped.startswith('```') or stripped.startswith('~~~'):
            fence = stripped[:3]
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].lstrip().startswith(fence):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # closing fence
            if lang.lower() == 'mermaid' and add_mermaid_diagram(doc, code_lines):
                continue
            add_code_block(doc, code_lines, lang)
            continue

        # Tables
        if is_table_start(lines, i):
            table_rows, new_idx = parse_table(lines, i)
            if table_rows:
                add_formatted_table(doc, table_rows)
            i = new_idx
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].lstrip().startswith('>'):
                qt = lines[i].lstrip()
                qt = qt[1:].lstrip() if qt.startswith('>') else qt
                quote_lines.append(qt.rstrip())
                i += 1
            quote_text = ' '.join(s for s in quote_lines if s)
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                add_left_border(p)
                parse_inline_formatting(p, quote_text)
                for run in p.runs:
                    if run.font.color.rgb is None:
                        run.font.color.rgb = COLOR_BLOCKQUOTE
                    run.italic = True
            continue

        # Lists
        if is_list_line(line):
            items, new_idx = parse_list_block(lines, i)
            if items:
                add_list_items(doc, items)
            i = new_idx
            continue

        # Paragraph: gather contiguous non-block lines
        para_lines = []
        while i < len(lines):
            cur = lines[i]
            cur_strip = cur.strip()
            if not cur_strip:
                break
            if cur_strip.startswith('#'):
                break
            if cur.lstrip().startswith('|') and is_table_start(lines, i):
                break
            if cur_strip.startswith('```') or cur_strip.startswith('~~~'):
                break
            if cur_strip.startswith('>'):
                break
            if is_list_line(cur):
                break
            if cur_strip in ('---', '***', '___'):
                break
            para_lines.append(cur.rstrip())
            i += 1

        if para_lines:
            # Markdown soft line break: trailing two spaces => hard break.
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            for n, pl in enumerate(para_lines):
                if n > 0:
                    if para_lines[n - 1].endswith('  '):
                        p.add_run().add_break()
                    else:
                        p.add_run(' ')
                parse_inline_formatting(p, pl.strip())

    doc.save(docx_file)
    return docx_file


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to Word (DOCX) with professional formatting',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python convert_md_to_docx.py document.md
  python convert_md_to_docx.py document.md output.docx
        """
    )
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('output', nargs='?', help='Output DOCX file (default: same name with .docx)')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found.")
        sys.exit(1)

    output_file = convert_md_to_docx(args.input, args.output)
    try:
        print(f"[OK] Created: {output_file}")
    except UnicodeEncodeError:
        sys.stdout.buffer.write(f"[OK] Created: {output_file}\n".encode('utf-8'))


if __name__ == '__main__':
    main()
